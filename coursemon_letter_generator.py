#!/usr/bin/env python
# coding: utf-8

import streamlit as st
from docx import Document
from docx.shared import Inches
from datetime import date
import openai

# -------------------
# App Setup
# -------------------
st.image("coursemon_pic_logo.jpg", width=150)
st.title("📄 Coursemon Letter Generator")

# Use API key from secrets
if "openai" not in st.secrets:
    st.error("❌ OpenAI API Key not found in secrets.")
    st.stop()

openai.api_key = st.secrets["openai"]["api_key"]

# -------------------
# Inputs
# -------------------
letter_type = st.selectbox("Letter Type", ["Employment Letter", "Business Proposal", "Notice Letter"])
name = st.text_input("Recipient Name", "Malaika Khan")
email = st.text_input("Recipient Email", "malaika@example.com")
position = st.text_input("Position/Role", "Business Development Specialist")

include_equity = False
equity = ""

if letter_type == "Employment Letter":
    include_equity = st.radio("Include Equity Offer?", ["Yes", "No"]) == "Yes"
    if include_equity:
        equity = st.text_input("Equity %", "5")

letter_date = st.date_input("Date", date.today())

# -------------------
# GPT Letter Generator
# -------------------
def get_letter_from_gpt(letter_type, name, position, equity, letter_date, include_equity):
    equity_line = f"\nInclude equity offering of {equity}% as part of this role." if include_equity and equity else ""
    
    prompt = f"""
You are a professional business assistant writing a formal {letter_type.lower()} on behalf of Ammar Jamshed, Founder & CEO of Coursemon.

Recipient: {name}
Position: {position}
Date: {letter_date.strftime('%B %d, %Y')}
{equity_line}

Write a structured letter with:
- A subject line
- Professional greeting
- Context and purpose
- Sections (such as Overview, Role Expectations, or Terms depending on the letter type)
- Closing with contact and sign-off:
Ammar Jamshed
Founder & CEO, Coursemon
contact@coursemon.com

Use a warm and professional tone suitable for startup and business communication.
    """

    response = openai.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that writes formal letters."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7,
        max_tokens=800
    )

    return response.choices[0].message.content

# -------------------
# Generate Letter
# -------------------
if st.button("Generate Letter with GPT"):
    with st.spinner("Generating letter with GPT..."):
        letter_text = get_letter_from_gpt(letter_type, name, position, equity, letter_date, include_equity)

        doc = Document()

        # Simulated Header
        header_table = doc.add_table(rows=1, cols=2)
        hdr_cells = header_table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run().add_picture("coursemon_pic_logo.jpg", width=Inches(1.2))
        hdr_cells[1].text = "Coursemon\nEmpowering Learning Journeys"
        header_table.style = 'Table Grid'

        doc.add_paragraph()  # Spacer

        doc.add_paragraph(f"Date: {letter_date.strftime('%B %d, %Y')}")
        doc.add_paragraph(f"To:\n{name}\n{email}")
        doc.add_paragraph()
        doc.add_paragraph(f"Subject: {letter_type} for {position}", style='Heading 2')
        doc.add_paragraph(letter_text)
        doc.add_paragraph()  # Spacer

        # Simulated Footer
        doc.add_paragraph("Warm regards,")
        doc.add_paragraph("Ammar Jamshed")
        doc.add_paragraph("Co-Founder & CEO, Coursemon")
        doc.add_paragraph("WhatsApp: 03412917004")
        doc.add_paragraph("https://coursemon.net")
        doc.add_paragraph("NED University, University Road, NIC Karachi")


        file_name = f"{name.replace(' ', '_')}_{letter_type.replace(' ', '_')}_Coursemon_Letter.docx"
        doc.save(file_name)

        with open(file_name, "rb") as f:
            st.success("✅ Letter Ready!")
            st.download_button(
                label="📥 Download Letter",
                data=f,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
